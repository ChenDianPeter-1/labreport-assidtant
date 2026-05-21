#!/usr/bin/env python3
"""Extract final figures and tables from postlab_report.html into a DOCX.

Requires python-docx and BeautifulSoup when run. The generated document still
must be render-checked and repaired if layout problems are visible.
"""

from __future__ import annotations

import argparse
from pathlib import Path


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--postlab", required=True, help="reports/postlab_report.html")
    parser.add_argument("--out", required=True, help="reports/appendix_figures_tables.docx")
    args = parser.parse_args()

    try:
        from bs4 import BeautifulSoup
        from docx import Document
        from docx.shared import Inches
    except Exception as exc:  # pragma: no cover - dependency guard
        raise SystemExit(
            "Missing dependency. Install/use an environment with beautifulsoup4 "
            f"and python-docx. Original error: {exc}"
        )

    postlab = Path(args.postlab)
    out = Path(args.out)
    soup = BeautifulSoup(postlab.read_text(encoding="utf-8"), "html.parser")
    doc = Document()
    doc.add_heading("附图附表", level=1)

    for node in soup.find_all(["table", "figure", "img"]):
        if node.name == "table":
            caption = ""
            prev = node.find_previous(["p", "h3", "h4", "caption"])
            if prev:
                caption = prev.get_text(" ", strip=True)
            if caption:
                doc.add_paragraph(caption)
            rows = node.find_all("tr")
            if rows:
                width = max(len(r.find_all(["th", "td"])) for r in rows)
                table = doc.add_table(rows=len(rows), cols=width)
                table.style = "Table Grid"
                for i, row in enumerate(rows):
                    cells = row.find_all(["th", "td"])
                    for j, cell in enumerate(cells):
                        table.cell(i, j).text = cell.get_text(" ", strip=True)
            doc.add_paragraph("")
        elif node.name in {"figure", "img"}:
            img = node.find("img") if node.name == "figure" else node
            if not img or not img.get("src"):
                continue
            img_path = (postlab.parent / img["src"]).resolve()
            if not img_path.exists():
                img_path = (postlab.parent.parent / img["src"]).resolve()
            caption_node = node.find("figcaption") if node.name == "figure" else None
            caption = caption_node.get_text(" ", strip=True) if caption_node else img.get("alt", "")
            if img_path.exists():
                doc.add_picture(str(img_path), width=Inches(5.6))
                if caption:
                    doc.add_paragraph(caption)
                doc.add_paragraph("")

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()

